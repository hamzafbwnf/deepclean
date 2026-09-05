from pathlib import Path
import re, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root=Path(__file__).parent
template=Path(r'C:\Users\hamza\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-design-report\assets\reference.docx')
doc=Document(template)
first_section=copy.deepcopy(doc.sections[0]._sectPr)
body=doc._element.body
for child in list(body):
    if child.tag != qn('w:sectPr'): body.remove(child)
normal=doc.styles['normal']
normal.font.name='Calibri'; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6)
normal.paragraph_format.line_spacing=1.08
for name in ['Title','Heading 1','Heading 2','Heading 3','Heading 4']:
    if name in doc.styles:
        doc.styles[name].font.color.rgb=RGBColor(0,0,0)
        doc.styles[name].paragraph_format.keep_with_next=True

def inline(p,text):
    for bit in re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)',text):
        if not bit: continue
        r=p.add_run(bit[2:-2] if bit.startswith('**') else bit[1:-1] if bit.startswith(('`','*')) else bit)
        if bit.startswith('**'):r.bold=True
        elif bit.startswith('`'):r.font.name='Consolas';r.font.size=Pt(10)
        elif bit.startswith('*'):r.italic=True
    return p

doc.add_paragraph('DEEPCLEAN', 'Title')
doc.add_paragraph('Oyun Tasarım Belgesi\nSürüm 0.7 • 5 Eylül 2026')
doc.add_picture(str(root/'Art/DEEPCLEAN_Boat_Modeling_Reference_01.png'),width=Inches(6.5))
doc.add_paragraph('Ekip paylaşım kopyası • Tam GDD ve tekne modelleme referansı')
doc.add_paragraph('Görsel genel modelleme yönünü gösterir; ölçülü teknik çizim değildir. Modelleme kapsamı ve açık noktalar: Bölüm 12.6.')
p=doc.add_paragraph();p._p.get_or_add_pPr().append(first_section)
doc.add_heading('Belgeyi kullanma rehberi',1)
doc.add_paragraph('Bu paylaşım kopyası, yaşayan GDD’nin tam içeriğini içerir. Karar durumları korunmuştur: aday fikirler ve açık sorular, uygulanması kesinleşmiş özellikler değildir.')
doc.add_paragraph('Modelleme için önce Bölüm 12.6’yı ve görseli, ardından Bölüm 16’daki seviye ve mini oyun bağlantılarını okuyun. Mevcut tekne gövdesini koruyun; yeni ekipmanları ayrı parçalar olarak hazırlayın. Uzman ve saha mühendisi düzeneklerinin ince bağlantıları henüz konsepttir.')
doc.add_paragraph('Bu dosya 5 Eylül 2026 tarihli paylaşım kopyasıdır. Ana tasarım kaynağı DEEPCLEAN_GDD.md dosyasıdır; değişiklikler ana GDD bakım görüşmesinde birleştirilir.')
doc.add_heading('İçerik',2)
source=(root/'DEEPCLEAN_GDD.md').read_text(encoding='utf-8-sig')
for line in source.splitlines():
    if line.startswith('## '):doc.add_paragraph(line[3:])
doc.add_page_break()
lines=source.splitlines();i=0
while i<len(lines):
    s=lines[i].strip();i+=1
    if not s or s=='---':continue
    if s.startswith('# '):continue
    if s.startswith('!['):
        m=re.match(r'!\[(.*?)\]\((.*?)\)',s)
        if m:
            doc.add_picture(str(root/m[2]),width=Inches(6.5))
            p=doc.add_paragraph(m[1]);p.runs[0].italic=True;p.runs[0].font.size=Pt(9)
        continue
    if s.startswith('|'):
        rows=[s]
        while i<len(lines) and lines[i].strip().startswith('|'):
            rows.append(lines[i].strip());i+=1
        vals=[[c.strip() for c in r.strip('|').split('|')] for r in rows if not re.match(r'^\|[\s:|\-]+\|$',r)]
        if not vals:continue
        n=len(vals[0]); t=doc.add_table(rows=0,cols=n);t.autofit=False
        lengths=[max(len(row[c]) if c<len(row) else 0 for row in vals) for c in range(n)]
        weights=[max(18,min(65,x**.7*4)) for x in lengths]
        widths=[6.5*w/sum(weights) for w in weights]
        for c,w in zip(t.columns,widths):c.width=Inches(w)
        for j,row in enumerate(vals):
            cells=t.add_row().cells
            for c,cell in enumerate(cells):
                cell.width=Inches(widths[c]);inline(cell.paragraphs[0],row[c] if c<len(row) else '')
                for p in cell.paragraphs:
                    p.paragraph_format.space_after=Pt(4);p.paragraph_format.space_before=Pt(4)
                    for r in p.runs:r.font.size=Pt(10);r.bold=True if j==0 else r.bold
                pr=cell._tc.get_or_add_tcPr();mar=OxmlElement('w:tcMar')
                for edge in ['top','left','bottom','right']:
                    e=OxmlElement('w:'+edge);e.set(qn('w:w'),'85');e.set(qn('w:type'),'dxa');mar.append(e)
                pr.append(mar)
            if j==0:
                h=OxmlElement('w:tblHeader');t.rows[j]._tr.get_or_add_trPr().append(h)
        borders=OxmlElement('w:tblBorders')
        for edge in ['top','left','bottom','right','insideH','insideV']:
            e=OxmlElement('w:'+edge);e.set(qn('w:val'),'single');e.set(qn('w:sz'),'4');e.set(qn('w:color'),'D9D9D9');borders.append(e)
        t._tbl.tblPr.append(borders);doc.add_paragraph()
        continue
    if s.startswith('##'):
        m=re.match(r'^(#+)\s+(.*)',s);level=min(len(m[1])-1,4)
        p=doc.add_heading('',level);inline(p,m[2]);continue
    if s.startswith('>'):s=s.lstrip('> ').rstrip()
    # Retain explicit numbering from the design, without restarting it globally.
    if s.startswith('- '):
        p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(.16);p.paragraph_format.first_line_indent=Inches(-.12);inline(p,'• '+s[2:])
    else:inline(doc.add_paragraph(),s)

out=root/'Share/DEEPCLEAN_GDD_v0.7_Ekip_Paylasim.docx'
out.parent.mkdir(exist_ok=True)
doc.save(out)
print(out)
print('paragraphs',len(doc.paragraphs),'tables',len(doc.tables),'images',len(doc.inline_shapes))
